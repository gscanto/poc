# main.py - Backend FastAPI para AutoReportAI (Versão Otimizada com Persistência)
import os
import json
import logging
import hashlib
import re
import pickle
from datetime import datetime
from typing import List, Optional, Dict
from pathlib import Path

from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import torch
from transformers import AutoTokenizer, AutoModelForCausalLM
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
from docx import Document
from docx.shared import Pt
from starlette.responses import StreamingResponse
import asyncio

# Configuração de logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Configuração da aplicação
app = FastAPI(title="AutoReportAI", version="2.0.0")

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configurações padrão
DEFAULT_CONFIG = {
    "style": "technical",
    "reference_format": "IEEE",
    "retrieve_references": True,
    "top_k": 5,  # Reduzido para evitar repetição
    "model_type": "local",  # Alterado para local
    "model_name": "microsoft/phi-2",  # Modelo padrão
    "min_words_per_section": 200,
    "max_reference_length": 500  # Limitar tamanho das referências
}

# Mapeamento de modelos disponíveis
AVAILABLE_MODELS = {
    "Phi-2 (2.7B parâmetros)": "microsoft/phi-2",
    "GPT-Neo-2.7B (2.7B parâmetros)": "EleutherAI/gpt-neo-2.7B",
    "GPT-J-6B (6B parâmetros)": "EleutherAI/gpt-j-6B",
    "Falcon-7B-Instruct (7B parâmetros)": "tiiuae/falcon-7b-instruct",
    "Llama-2-7B-Chat (7B parâmetros)": "meta-llama/Llama-2-7b-chat-hf",
    "Mistral-7B-Instruct (7B parâmetros)": "mistralai/Mistral-7B-Instruct-v0.1"
}

# Diretórios de persistência
PERSISTENCE_DIR = Path("persistence")
CORPUS_METADATA_FILE = PERSISTENCE_DIR / "corpus_metadata.json"
FAISS_INDEX_FILE = PERSISTENCE_DIR / "faiss_index.bin"

# Modelos Pydantic
class ReportRequest(BaseModel):
    title: str
    context: str
    sections: List[str]

class ReportResponse(BaseModel):
    report_id: str
    content: str
    references: List[dict]
    generation_time: float
    tokens_used: int

class DocumentProcessor:
    """Processa e limpa documentos antes de indexar"""
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Remove formatação desnecessária e limpa texto"""
        # Remove múltiplas quebras de linha
        text = re.sub(r'\n{3,}', '\n\n', text)
        # Remove linhas muito curtas (provavelmente títulos/índices)
        lines = text.split('\n')
        cleaned_lines = [line for line in lines if len(line.strip()) > 20 or line.strip() == '']
        text = '\n'.join(cleaned_lines)
        # Remove espaços múltiplos
        text = re.sub(r' {2,}', ' ', text)
        return text.strip()
    
    @staticmethod
    def extract_meaningful_chunks(text: str, chunk_size: int = 800) -> List[str]:
        """Extrai chunks significativos do texto"""
        # Divide por parágrafos
        paragraphs = [p.strip() for p in text.split('\n\n') if len(p.strip()) > 100]
        
        chunks = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) < chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    chunks.append(current_chunk.strip())
                current_chunk = para + "\n\n"
        
        if current_chunk:
            chunks.append(current_chunk.strip())
        
        return chunks[:10]  # Limitar a 10 chunks por documento

class PersistenceManager:
    """Gerencia persistência de índices e metadados"""
    
    def __init__(self, persistence_dir: Path):
        self.persistence_dir = persistence_dir
        self.persistence_dir.mkdir(parents=True, exist_ok=True)
        self.metadata_file = self.persistence_dir / "corpus_metadata.json"
        self.index_file = self.persistence_dir / "faiss_index.bin"
        self.backup_dir = self.persistence_dir / "backups"
        self.backup_dir.mkdir(parents=True, exist_ok=True)
    
    def save_corpus(self, corpus_metadata: List[dict], faiss_index: faiss.Index) -> bool:
        """Salva corpus e índice FAISS em disco"""
        try:
            # Backup antes de salvar
            self._create_backup()
            
            # Salvar metadados
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(corpus_metadata, f, ensure_ascii=False, indent=2)
            
            # Salvar índice FAISS
            faiss.write_index(faiss_index, str(self.index_file))
            
            logger.info(f"✓ Corpus persistido com sucesso: {len(corpus_metadata)} documentos")
            return True
            
        except Exception as e:
            logger.error(f"✗ Erro ao persistir corpus: {e}")
            return False
    
    def load_corpus(self) -> tuple[List[dict], Optional[faiss.Index]]:
        """Carrega corpus e índice FAISS do disco"""
        try:
            # Carregar metadados
            if not self.metadata_file.exists():
                logger.info("Nenhum corpus persistido encontrado")
                return [], None
            
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                corpus_metadata = json.load(f)
            
            # Carregar índice FAISS
            if not self.index_file.exists():
                logger.warning("Metadados encontrados mas índice FAISS ausente")
                return [], None
            
            faiss_index = faiss.read_index(str(self.index_file))
            
            logger.info(f"✓ Corpus carregado: {len(corpus_metadata)} documentos, {faiss_index.ntotal} vetores")
            return corpus_metadata, faiss_index
            
        except Exception as e:
            logger.error(f"✗ Erro ao carregar corpus: {e}")
            return [], None
    
    def _create_backup(self):
        """Cria backup dos arquivos existentes"""
        try:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            if self.metadata_file.exists():
                backup_metadata = self.backup_dir / f"corpus_metadata_{timestamp}.json"
                with open(self.metadata_file, 'r') as src, open(backup_metadata, 'w') as dst:
                    dst.write(src.read())
            
            if self.index_file.exists():
                backup_index = self.backup_dir / f"faiss_index_{timestamp}.bin"
                faiss_index = faiss.read_index(str(self.index_file))
                faiss.write_index(faiss_index, str(backup_index))
            
            # Manter apenas últimos 5 backups
            self._cleanup_old_backups(keep=5)
            
        except Exception as e:
            logger.warning(f"Erro ao criar backup: {e}")
    
    def _cleanup_old_backups(self, keep: int = 5):
        """Remove backups antigos, mantendo apenas os mais recentes"""
        try:
            metadata_backups = sorted(self.backup_dir.glob("corpus_metadata_*.json"))
            index_backups = sorted(self.backup_dir.glob("faiss_index_*.bin"))
            
            for backup in metadata_backups[:-keep]:
                backup.unlink()
            
            for backup in index_backups[:-keep]:
                backup.unlink()
                
        except Exception as e:
            logger.warning(f"Erro ao limpar backups: {e}")
    
    def get_stats(self) -> dict:
        """Retorna estatísticas de persistência"""
        stats = {
            "metadata_exists": self.metadata_file.exists(),
            "index_exists": self.index_file.exists(),
            "metadata_size_kb": 0,
            "index_size_kb": 0,
            "last_modified": None,
            "num_backups": 0
        }
        
        if self.metadata_file.exists():
            stats["metadata_size_kb"] = round(self.metadata_file.stat().st_size / 1024, 2)
            stats["last_modified"] = datetime.fromtimestamp(
                self.metadata_file.stat().st_mtime
            ).strftime("%Y-%m-%d %H:%M:%S")
        
        if self.index_file.exists():
            stats["index_size_kb"] = round(self.index_file.stat().st_size / 1024, 2)
        
        stats["num_backups"] = len(list(self.backup_dir.glob("corpus_metadata_*.json")))
        
        return stats

class ModelManager:
    def __init__(self):
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.embedding_model = None
        self.llm_model = None
        self.tokenizer = None
        self.faiss_index = None
        self.corpus_metadata = []
        self.model_name = None
        self.report_cache = {}
        self.document_processor = DocumentProcessor()
        self.persistence_manager = PersistenceManager(PERSISTENCE_DIR)
        
        logger.info(f"Device: {self.device}")

    def _generate_cache_key(self, context: str, sections: List[str]) -> str:
        """Gera chave de cache baseada no contexto e seções"""
        sorted_sections = sorted(sections)
        combined = context + "|" + "|".join(sorted_sections)
        return hashlib.sha256(combined.encode('utf-8')).hexdigest()
    
    # ADICIONE ESTES MÉTODOS À CLASSE ModelManager NO main.py
# Insira após o método _generate_cache_key e antes do método load_embedding_model

    def _format_references_for_prompt(self, docs: List[dict], max_length: int = 1000) -> str:
        """Formata referências para inclusão no prompt do modelo"""
        if not docs:
            return "Nenhuma referência disponível."
        
        references_text = []
        total_length = 0
        
        for i, doc in enumerate(docs, 1):
            text = doc.get('text', '')
            title = doc.get('original_title', doc.get('title', 'Untitled'))
            
            # Limitar tamanho de cada referência
            max_ref_length = max_length // len(docs)
            if len(text) > max_ref_length:
                text = text[:max_ref_length] + "..."
            
            ref_text = f"\n[Fonte {i}] {title}:\n{text}\n"
            
            if total_length + len(ref_text) > max_length:
                break
            
            references_text.append(ref_text)
            total_length += len(ref_text)
        
        return "\n".join(references_text)

    def load_embedding_model(self, model_name="all-MiniLM-L6-v2"):
        """Carrega modelo de embeddings"""
        if self.embedding_model is not None:
            return
        logger.info(f"Carregando modelo de embeddings: {model_name}")
        self.embedding_model = SentenceTransformer(model_name)
        if self.device == "cuda":
            try:
                self.embedding_model = self.embedding_model.to(self.device)
            except Exception:
                logger.warning("Executando embedding em CPU.")
        logger.info("Modelo de embeddings carregado com sucesso")

    def load_language_model(self, model_name="microsoft/phi-2"):
        """Carrega modelo de linguagem para geração de texto"""
        if self.llm_model is not None and self.model_name == model_name:
            logger.info(f"Modelo {model_name} já carregado")
            return
        
        try:
            logger.info(f"Carregando modelo de linguagem: {model_name}")
            
            # Liberar modelo anterior se existir
            if self.llm_model is not None:
                del self.llm_model
                del self.tokenizer
                if torch.cuda.is_available():
                    torch.cuda.empty_cache()
            
            # Carregar tokenizer
            self.tokenizer = AutoTokenizer.from_pretrained(
                model_name,
                trust_remote_code=True
            )
            
            # Configurar padding token se não existir
            if self.tokenizer.pad_token is None:
                self.tokenizer.pad_token = self.tokenizer.eos_token
            
            # Carregar modelo
            self.llm_model = AutoModelForCausalLM.from_pretrained(
                model_name,
                torch_dtype=torch.float16 if self.device == "cuda" else torch.float32,
                device_map="auto" if self.device == "cuda" else None,
                trust_remote_code=True,
                low_cpu_mem_usage=True
            )
            
            if self.device == "cpu":
                self.llm_model = self.llm_model.to(self.device)
            
            self.llm_model.eval()
            self.model_name = model_name
            
            logger.info(f"Modelo {model_name} carregado com sucesso em {self.device}")
            
        except Exception as e:
            logger.error(f"Erro ao carregar modelo {model_name}: {e}")
            logger.info("Usando geração por template como fallback")
            self.llm_model = None
            self.tokenizer = None
            self.model_name = None

    def create_faiss_index(self, dimension=384):
        """Cria índice FAISS"""
        if self.faiss_index is None:
            self.faiss_index = faiss.IndexFlatIP(dimension)
            logger.info("Índice FAISS criado")

    def load_persisted_corpus(self) -> bool:
        """Carrega corpus persistido do disco"""
        corpus_metadata, faiss_index = self.persistence_manager.load_corpus()
        
        if corpus_metadata and faiss_index:
            self.corpus_metadata = corpus_metadata
            self.faiss_index = faiss_index
            logger.info(f"✓ Corpus persistido restaurado: {len(corpus_metadata)} documentos")
            return True
        
        return False

    def save_corpus(self) -> bool:
        """Salva corpus atual em disco"""
        if not self.corpus_metadata or not self.faiss_index:
            logger.warning("Nenhum corpus para persistir")
            return False
        
        return self.persistence_manager.save_corpus(self.corpus_metadata, self.faiss_index)

    def add_documents_to_index(self, documents: List[dict], auto_save: bool = True):
        """Adiciona documentos ao índice FAISS com processamento inteligente e persistência automática"""
        if self.embedding_model is None:
            self.load_embedding_model()

        processed_docs = []
        
        for doc in documents:
            original_text = doc.get('text', '')
            if not original_text or len(original_text) < 100:
                continue
            
            # Limpar texto
            cleaned_text = self.document_processor.clean_text(original_text)
            
            # Extrair chunks significativos
            chunks = self.document_processor.extract_meaningful_chunks(cleaned_text)
            
            # Criar um documento por chunk
            for i, chunk in enumerate(chunks):
                chunk_doc = doc.copy()
                chunk_doc['text'] = chunk
                chunk_doc['chunk_id'] = i
                chunk_doc['original_title'] = doc.get('title', 'Untitled')
                chunk_doc['upload_timestamp'] = datetime.now().isoformat()
                processed_docs.append(chunk_doc)
        
        if not processed_docs:
            logger.warning("Nenhum documento válido para indexar")
            return
        
        texts = [doc['text'] for doc in processed_docs]
        embeddings = self.embedding_model.encode(texts, convert_to_numpy=True, show_progress_bar=True)
        
        if embeddings is None or len(embeddings) == 0:
            logger.warning("Nenhum embedding foi gerado.")
            return

        faiss.normalize_L2(embeddings)
        self.faiss_index.add(embeddings)
        self.corpus_metadata.extend(processed_docs)
        logger.info(f"{len(processed_docs)} chunks adicionados de {len(documents)} documentos (total: {self.faiss_index.ntotal})")
        
        # Persistência automática
        if auto_save:
            success = self.save_corpus()
            if success:
                logger.info("✓ Corpus automaticamente persistido")
            else:
                logger.warning("✗ Falha na persistência automática")

    def retrieve_documents(self, query: str, top_k: int = 5) -> List[dict]:
        """Recupera documentos relevantes evitando duplicatas"""
        if (self.faiss_index is None) or (self.faiss_index.ntotal == 0):
            logger.warning("Índice FAISS vazio")
            return []

        query_embedding = self.embedding_model.encode([query], convert_to_numpy=True)
        if query_embedding is None or len(query_embedding) == 0:
            logger.warning("Embedding da query retornou None")
            return []

        faiss.normalize_L2(query_embedding)
        # Buscar mais resultados para filtrar duplicatas
        k = min(top_k * 3, self.faiss_index.ntotal)
        distances, indices = self.faiss_index.search(query_embedding, k)

        results = []
        seen_titles = set()
        
        for idx, score in zip(indices[0], distances[0]):
            if idx < len(self.corpus_metadata):
                doc = self.corpus_metadata[idx].copy()
                doc["score"] = float(score)
                
                # Evitar documentos duplicados do mesmo arquivo
                title = doc.get('original_title', doc.get('title', ''))
                if title not in seen_titles:
                    results.append(doc)
                    seen_titles.add(title)
                    
                if len(results) >= top_k:
                    break
        
        logger.info(f"Recuperados {len(results)} documentos únicos para query")
        return results

    def generate_section(self, section_name: str, context: str, retrieved_docs: List[dict], config: dict):
        """Gera conteúdo da seção com qualidade aprimorada usando modelo local"""
        style = config.get('style', 'technical')
        model_type = config.get('model_type', 'local')
        model_name = config.get('model_name', 'microsoft/phi-2')
        min_words = config.get('min_words_per_section', 200)

        logger.info(f"Gerando seção '{section_name}' usando modelo local: {model_name}")

        # Sempre usar modelo local
        content, tokens = self._generate_section_local(section_name, context, retrieved_docs, style, model_name)
        
        # Pós-processamento: remover duplicações e melhorar formatação
        content = self._post_process_content(content, section_name)
        
        # Validar qualidade mínima
        word_count = len(content.split())
        if word_count < min_words * 0.7:  # 70% do mínimo
            logger.warning(f"Conteúdo muito curto ({word_count} palavras), regenerando...")
            content = self._generate_enhanced_fallback(section_name, context, retrieved_docs, min_words)
            tokens = 0
        
        return content, tokens

    def _post_process_content(self, content: str, section_name: str) -> str:
        """Remove duplicações e melhora formatação"""
        # Remover título duplicado
        lines = content.split('\n')
        cleaned_lines = []
        seen_lines = set()
        
        for line in lines:
            line_clean = line.strip()
            # Remover linhas duplicadas
            if line_clean and line_clean not in seen_lines:
                cleaned_lines.append(line)
                seen_lines.add(line_clean)
            elif not line_clean:
                cleaned_lines.append(line)
        
        content = '\n'.join(cleaned_lines)
        
        # Remover múltiplos títulos da seção
        content = re.sub(rf'#{1,3}\s*{re.escape(section_name)}\s*\n+', '', content, flags=re.IGNORECASE)
        
        # Limpar referências mal formatadas (Source X repetido)
        content = re.sub(r'(\[Source \d+\])\s*\1+', r'\1', content)
        
        return content.strip()

    def _generate_section_local(self, section_name: str, context: str, retrieved_docs: List[dict], 
                                  style: str, model_name: str = "microsoft/phi-2"):
        """Gera conteúdo usando modelo de linguagem local"""
        
        # Carregar modelo se necessário
        if self.llm_model is None or self.model_name != model_name:
            self.load_language_model(model_name)
        
        # Se o modelo não carregou, usar fallback
        if self.llm_model is None:
            logger.warning("Modelo não disponível, usando enhanced fallback")
            return self._generate_enhanced_fallback(section_name, context, retrieved_docs, 200), 0
        
        try:
            # Preparar contexto das referências
            references_context = self._format_references_for_prompt(retrieved_docs, max_length=1000)
            
            # Construir prompt otimizado para modelo local
            prompt = self._build_local_model_prompt(section_name, context, references_context, style)
            
            # Tokenizar
            inputs = self.tokenizer(
                prompt,
                return_tensors="pt",
                max_length=2048,
                truncation=True,
                padding=True
            )
            
            inputs = {k: v.to(self.device) for k, v in inputs.items()}
            
            # Gerar texto
            logger.info(f"Gerando texto com {model_name}...")
            with torch.no_grad():
                outputs = self.llm_model.generate(
                    **inputs,
                    max_new_tokens=800,
                    temperature=0.7,
                    top_p=0.9,
                    do_sample=True,
                    pad_token_id=self.tokenizer.pad_token_id,
                    eos_token_id=self.tokenizer.eos_token_id,
                    repetition_penalty=1.2,
                    no_repeat_ngram_size=3
                )
            
            # Decodificar
            generated_text = self.tokenizer.decode(outputs[0], skip_special_tokens=True)
            
            # Remover o prompt da resposta
            if prompt in generated_text:
                generated_text = generated_text.replace(prompt, "").strip()
            
            # Limpar texto gerado
            generated_text = self._clean_generated_text(generated_text)
            
            word_count = len(generated_text.split())
            tokens_used = len(outputs[0])
            
            logger.info(f"Seção gerada com sucesso: {word_count} palavras, {tokens_used} tokens")
            
            return generated_text, tokens_used
            
        except Exception as e:
            logger.error(f"Erro na geração com modelo local: {e}")
            logger.info("Usando enhanced fallback")
            return self._generate_enhanced_fallback(section_name, context, retrieved_docs, 200), 0
    
    def _build_local_model_prompt(self, section_name: str, context: str, 
                                   references: str, style: str) -> str:
        """Constrói prompt otimizado para modelos locais"""
        
        style_instructions = {
            "technical": "Use linguagem técnica e formal.",
            "concise": "Seja direto e objetivo.",
            "detailed": "Forneça explicações detalhadas."
        }
        style_guide = style_instructions.get(style, style_instructions["technical"])
        
        prompt = f"""Escreva a seção "{section_name}" de um relatório técnico.

Contexto: {context[:1000]}

Referências disponíveis:
{references}

Instruções:
- {style_guide}
- Escreva pelo menos 200 palavras
- Use parágrafos bem estruturados
- Cite as fontes usando [Fonte X]
- Seja informativo e preciso

Seção "{section_name}":
"""
        
        return prompt
    
    def _clean_generated_text(self, text: str) -> str:
        """Limpa texto gerado pelo modelo"""
        # Remover linhas muito curtas no início
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if len(line) > 20 or not line:
                cleaned_lines.append(line)
        
        text = '\n'.join(cleaned_lines)
        
        # Remover repetições excessivas
        text = re.sub(r'(\b\w+\b)(\s+\1){3,}', r'\1', text)
        
        # Limitar parágrafos muito longos
        paragraphs = text.split('\n\n')
        cleaned_paragraphs = []
        
        for para in paragraphs:
            if len(para.split()) > 200:  # Parágrafo muito longo
                sentences = para.split('. ')
                cleaned_paragraphs.append('. '.join(sentences[:8]) + '.')
            else:
                cleaned_paragraphs.append(para)
        
        return '\n\n'.join(cleaned_paragraphs).strip()

    def _generate_enhanced_fallback(self, section_name: str, context: str, 
                                     docs: List[dict], min_words: int = 200) -> str:
        """Gera conteúdo de fallback com qualidade superior"""
        paragraphs = []
        
        # Parágrafo 1: Introdução
        intro = f"A seção de {section_name} representa um componente fundamental deste estudo. "
        if context:
            intro += f"Considerando que {context[:200].lower()}, "
            intro += "torna-se essencial examinar detalhadamente os aspectos teóricos e práticos envolvidos. "
            intro += "Esta análise visa estabelecer uma compreensão abrangente do tema e suas implicações."
        paragraphs.append(intro)
        
        # Parágrafos 2-4: Desenvolvimento baseado em referências
        if docs:
            for i, doc in enumerate(docs[:3], 1):
                text = doc.get('text', '')
                if len(text) < 50:
                    continue
                
                # Extrair informação relevante
                sentences = [s.strip() for s in text.split('.') if len(s.strip()) > 30]
                relevant_info = '. '.join(sentences[:3])
                
                title = doc.get('original_title', doc.get('title', 'a literatura'))
                
                para = f"De acordo com {title} [Fonte {i}], {relevant_info.lower()}. "
                para += "Esta perspectiva oferece insights valiosos para a compreensão do fenômeno em questão. "
                para += "A análise destes elementos permite estabelecer conexões importantes entre teoria e prática, "
                para += "demonstrando a aplicabilidade dos conceitos discutidos em contextos reais."
                paragraphs.append(para)
        
        # Parágrafo de análise
        analysis = "A integração dos conceitos apresentados revela uma complexidade inerente ao tema. "
        analysis += "Diversos fatores interconectados contribuem para a dinâmica observada, "
        analysis += "incluindo aspectos técnicos, metodológicos e contextuais. "
        analysis += "A consideração holística destes elementos é crucial para uma compreensão adequada "
        analysis += "e para o desenvolvimento de abordagens eficazes."
        paragraphs.append(analysis)
        
        # Conclusão da seção
        conclusion = f"Em síntese, os aspectos discutidos em {section_name} estabelecem fundamentos sólidos "
        conclusion += "para o desenvolvimento subsequente deste trabalho. "
        conclusion += "As evidências apresentadas, suportadas pelas referências consultadas, "
        conclusion += "demonstram a relevância e aplicabilidade dos conceitos abordados. "
        conclusion += "Este entendimento fornece a base necessária para análises mais aprofundadas "
        conclusion += "e para a formulação de conclusões robustas e bem fundamentadas."
        paragraphs.append(conclusion)
        
        content = "\n\n".join(paragraphs)
        
        # Verificar se atingiu palavra mínima
        word_count = len(content.split())
        if word_count < min_words:
            # Adicionar parágrafo extra de expansão
            expansion = "Adicionalmente, é importante ressaltar que a evolução histórica deste campo "
            expansion += "demonstra uma progressão constante em direção a abordagens mais sofisticadas e eficazes. "
            expansion += "As contribuições de diversos pesquisadores e profissionais ao longo do tempo "
            expansion += "enriqueceram significativamente nosso entendimento, estabelecendo paradigmas "
            expansion += "que continuam a orientar investigações contemporâneas e futuras direções de pesquisa."
            paragraphs.insert(-1, expansion)
            content = "\n\n".join(paragraphs)
        
        return content

    def clear_corpus(self) -> bool:
        """Limpa todo o corpus e índice"""
        try:
            self.corpus_metadata = []
            if self.faiss_index:
                dimension = self.faiss_index.d
                self.faiss_index = faiss.IndexFlatIP(dimension)
            
            # Remover arquivos persistidos
            if CORPUS_METADATA_FILE.exists():
                CORPUS_METADATA_FILE.unlink()
            if FAISS_INDEX_FILE.exists():
                FAISS_INDEX_FILE.unlink()
            
            logger.info("✓ Corpus limpo com sucesso")
            return True
        except Exception as e:
            logger.error(f"✗ Erro ao limpar corpus: {e}")
            return False


# Instância global
model_manager = ModelManager()

# Rotas da API
@app.on_event("startup")
async def startup_event():
    """Inicialização com carregamento de corpus persistido e modelo local"""
    logger.info("=" * 60)
    logger.info("Inicializando AutoReportAI v2.0 com Persistência...")
    logger.info("=" * 60)
    
    try:
        # Carregar modelo de embeddings
        model_manager.load_embedding_model()
        model_manager.create_faiss_index()
        
        # Carregar modelo de linguagem padrão
        logger.info("Carregando modelo de linguagem padrão...")
        model_manager.load_language_model(DEFAULT_CONFIG["model_name"])
        
        # Tentar carregar corpus persistido
        corpus_loaded = model_manager.load_persisted_corpus()
        
        # Se não há corpus persistido, carregar corpus de exemplo
        if not corpus_loaded:
            logger.info("Nenhum corpus persistido encontrado. Carregando corpus de exemplo...")
            sample_corpus = load_sample_corpus()
            if sample_corpus:
                model_manager.add_documents_to_index(sample_corpus, auto_save=True)
        
        # Exibir estatísticas
        stats = model_manager.persistence_manager.get_stats()
        logger.info("=" * 60)
        logger.info("STATUS DA PERSISTÊNCIA:")
        logger.info(f"  Metadados: {'✓' if stats['metadata_exists'] else '✗'} ({stats['metadata_size_kb']} KB)")
        logger.info(f"  Índice FAISS: {'✓' if stats['index_exists'] else '✗'} ({stats['index_size_kb']} KB)")
        logger.info(f"  Última modificação: {stats['last_modified']}")
        logger.info(f"  Backups disponíveis: {stats['num_backups']}")
        logger.info(f"  Documentos indexados: {model_manager.faiss_index.ntotal if model_manager.faiss_index else 0}")
        logger.info("=" * 60)
        logger.info("MODELO DE LINGUAGEM:")
        logger.info(f"  Modelo ativo: {model_manager.model_name or 'Nenhum (usando fallback)'}")
        logger.info(f"  Device: {model_manager.device}")
        logger.info("=" * 60)
        logger.info("✓ AutoReportAI v2.0 inicializado com sucesso!")
        logger.info("=" * 60)
        
    except Exception as e:
        logger.error(f"✗ Erro na inicialização: {e}")

@app.get("/")
async def root():
    persistence_stats = model_manager.persistence_manager.get_stats()
    return {
        "name": "AutoReportAI",
        "version": "2.0.0",
        "status": "running",
        "device": model_manager.device,
        "documents_indexed": model_manager.faiss_index.ntotal if model_manager.faiss_index else 0,
        "llm_model": model_manager.model_name or "Fallback (template-based)",
        "model_loaded": model_manager.llm_model is not None,
        "persistence": {
            "enabled": True,
            "metadata_exists": persistence_stats["metadata_exists"],
            "index_exists": persistence_stats["index_exists"],
            "last_save": persistence_stats["last_modified"],
            "backups": persistence_stats["num_backups"]
        }
    }

@app.post("/generate-report", response_model=ReportResponse)
def generate_report(request: ReportRequest):
    """Gera relatório técnico de alta qualidade"""
    start_time = datetime.now()
    config = DEFAULT_CONFIG.copy()

    cache_key = model_manager._generate_cache_key(request.context, request.sections)

    if cache_key in model_manager.report_cache:
        cached_data = model_manager.report_cache[cache_key]
        logger.info(f"Relatório encontrado no cache: {request.title}")
        report_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_report(report_id, cached_data["content"], cached_data["references"])
        return ReportResponse(
            report_id=report_id,
            content=cached_data["content"],
            references=cached_data["references"],
            generation_time=0.0,
            tokens_used=0,
        )

    logger.info(f"Gerando relatório: {request.title}")

    try:
        retrieved_docs = model_manager.retrieve_documents(
            request.context,
            top_k=config['top_k']
        ) if config['retrieve_references'] else []

        logger.info(f"Documentos recuperados: {len(retrieved_docs)}")

        sections = []
        total_tokens = 0

        for i, sec in enumerate(request.sections, 1):
            logger.info(f"Gerando seção {i}/{len(request.sections)}: {sec}")
            content, tokens = model_manager.generate_section(
                sec,
                request.context,
                retrieved_docs,
                config
            )
            word_count = len(content.split())
            logger.info(f"Seção '{sec}' gerada: {word_count} palavras, {len(content)} caracteres")
            
            sections.append({"title": sec, "content": content})
            total_tokens += tokens

        content_md = format_report(
            request.title,
            sections,
            retrieved_docs,
            config['reference_format']
        )

        model_manager.report_cache[cache_key] = {
            "content": content_md,
            "references": retrieved_docs
        }

        report_id = datetime.now().strftime("%Y%m%d_%H%M%S")
        save_report(report_id, content_md, retrieved_docs)

        generation_time = (datetime.now() - start_time).total_seconds()
        logger.info(f"Relatório gerado com sucesso em {generation_time:.2f}s")

        return ReportResponse(
            report_id=report_id,
            content=content_md,
            references=retrieved_docs,
            generation_time=generation_time,
            tokens_used=total_tokens,
        )

    except Exception as e:
        logger.exception("Erro inesperado ao gerar relatório.")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-documents")
async def upload_documents(documents: List[dict]):
    """Adiciona documentos ao corpus com persistência automática"""
    try:
        initial_count = model_manager.faiss_index.ntotal if model_manager.faiss_index else 0
        
        model_manager.add_documents_to_index(documents, auto_save=True)
        
        final_count = model_manager.faiss_index.ntotal if model_manager.faiss_index else 0
        added = final_count - initial_count
        
        return {
            "status": "success",
            "documents_uploaded": len(documents),
            "vectors_added": added,
            "total_vectors": final_count,
            "persisted": True,
            "message": f"✓ {len(documents)} documentos processados e {added} vetores adicionados (persistido automaticamente)"
        }
    except Exception as e:
        logger.error(f"Erro ao fazer upload de documentos: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/save-corpus")
async def save_corpus_endpoint():
    """Salva manualmente o corpus atual"""
    try:
        success = model_manager.save_corpus()
        if success:
            stats = model_manager.persistence_manager.get_stats()
            return {
                "status": "success",
                "message": "Corpus salvo com sucesso",
                "stats": stats
            }
        else:
            raise HTTPException(status_code=500, detail="Falha ao salvar corpus")
    except Exception as e:
        logger.error(f"Erro ao salvar corpus: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/change-model")
async def change_model(model_name: str):
    """Troca o modelo de linguagem ativo"""
    try:
        # Verificar se o modelo está na lista de disponíveis
        available = list(AVAILABLE_MODELS.values())
        if model_name not in available:
            raise HTTPException(status_code=400, detail=f"Modelo {model_name} não disponível")
        
        logger.info(f"Trocando modelo para: {model_name}")
        model_manager.load_language_model(model_name)
        
        return {
            "status": "success",
            "model_name": model_name,
            "model_loaded": model_manager.model_name == model_name,
            "device": model_manager.device
        }
    except Exception as e:
        logger.error(f"Erro ao trocar modelo: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/current-model")
async def get_current_model():
    """Retorna informações sobre o modelo atual"""
    return {
        "model_name": model_manager.model_name or "Nenhum (usando fallback)",
        "model_loaded": model_manager.llm_model is not None,
        "device": model_manager.device,
        "available_models": AVAILABLE_MODELS
    }

@app.post("/clear-corpus")
async def clear_corpus_endpoint():
    """Limpa todo o corpus e arquivos persistidos"""
    try:
        success = model_manager.clear_corpus()
        if success:
            return {
                "status": "success",
                "message": "Corpus limpo com sucesso",
                "documents_remaining": 0
            }
        else:
            raise HTTPException(status_code=500, detail="Falha ao limpar corpus")
    except Exception as e:
        logger.error(f"Erro ao limpar corpus: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/persistence-stats")
async def persistence_stats():
    """Retorna estatísticas detalhadas de persistência"""
    stats = model_manager.persistence_manager.get_stats()
    stats["documents_in_memory"] = len(model_manager.corpus_metadata)
    stats["vectors_indexed"] = model_manager.faiss_index.ntotal if model_manager.faiss_index else 0
    return stats

@app.get("/corpus-stats")
async def corpus_stats():
    return {
        "total_documents": model_manager.faiss_index.ntotal if model_manager.faiss_index else 0,
        "unique_documents": len(set(
            doc.get('original_title', doc.get('title', ''))
            for doc in model_manager.corpus_metadata
        )),
        "embedding_dimension": 384,
        "device": model_manager.device,
        "persisted": model_manager.persistence_manager.get_stats()["metadata_exists"]
    }

@app.get("/corpus-documents")
async def corpus_documents():
    documents = []
    seen_titles = set()
    for doc in model_manager.corpus_metadata:
        title = doc.get('original_title', doc.get('title', 'Untitled'))
        if title not in seen_titles:
            # Calculate total text length for all chunks of this document
            total_text_length = sum(len(d.get('text', '')) for d in model_manager.corpus_metadata
                                  if d.get('original_title') == title)
            
            upload_time = doc.get('upload_timestamp', 'Unknown')
            if upload_time != 'Unknown':
                try:
                    upload_time = datetime.fromisoformat(upload_time).strftime("%Y-%m-%d %H:%M:%S")
                except:
                    pass
            
            documents.append({
                "id": doc.get("id", ""),
                "title": title,
                "source": doc.get("source", "Unknown"),
                "year": doc.get("year", ""),
                "text_length": total_text_length,
                "chunks": sum(1 for d in model_manager.corpus_metadata if d.get('original_title') == title),
                "uploaded_at": upload_time
            })
            seen_titles.add(title)
    
    return {
        "documents": documents,
        "total_unique_documents": len(documents),
        "total_chunks": len(model_manager.corpus_metadata)
    }

@app.get("/reports")
async def list_reports():
    reports_dir = Path("reports")
    reports = []
    if reports_dir.exists():
        for md_file in reports_dir.glob("*.md"):
            report_id = md_file.stem
            stat = md_file.stat()
            size_kb = stat.st_size / 1024
            mtime = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
            reports.append({
                "id": report_id,
                "title": f"Relatório {report_id}",
                "generated_at": mtime,
                "size_kb": round(size_kb, 1)
            })
    return {"reports": reports}

def load_sample_corpus():
    """Corpus de exemplo de qualidade"""
    return [
        {
            "id": "ml_intro",
            "title": "Fundamentos de Machine Learning",
            "text": """Machine learning representa um paradigma revolucionário na ciência da computação, permitindo que sistemas computacionais aprendam e melhorem seu desempenho através da experiência, sem serem explicitamente programados para cada tarefa específica.

A essência do machine learning reside na capacidade de identificar padrões em dados e utilizar esses padrões para fazer previsões ou tomar decisões. Este processo envolve três componentes fundamentais: dados de treinamento, um algoritmo de aprendizado e um modelo resultante.

Os algoritmos de machine learning podem ser categorizados em três paradigmas principais. O aprendizado supervisionado utiliza dados rotulados para treinar modelos que mapeiam entradas para saídas conhecidas. O aprendizado não supervisionado descobre estruturas ocultas em dados não rotulados. O aprendizado por reforço aprende através de interação com um ambiente, recebendo recompensas ou penalidades.

Aplicações práticas de machine learning permeiam diversos setores: reconhecimento de imagem e voz, sistemas de recomendação, diagnóstico médico, previsão financeira, veículos autônomos e processamento de linguagem natural.""",
            "source": "Journal of Artificial Intelligence Research",
            "year": "2024",
            "authors": ["Silva, M.", "Santos, R."]
        }
    ]

def format_report(title: str, sections: List[dict], references: List[dict], ref_format: str) -> str:
    """Formata relatório em Markdown de alta qualidade"""
    date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    report_lines = []
    
    report_lines.append(f"# {title}\n")
    report_lines.append(f"**Data de Geração:** {date_str}\n")
    report_lines.append("---\n")

    report_lines.append("## Sumário\n")
    for i, sec in enumerate(sections, start=1):
        report_lines.append(f"{i}. {sec['title']}")
    report_lines.append("\n---\n")

    for sec in sections:
        report_lines.append(f"\n## {sec['title']}\n")
        report_lines.append(sec['content'])
        report_lines.append("\n")

    report_lines.append("\n---\n")
    report_lines.append("## Referências\n")
    
    if references:
        # Remover duplicatas
        seen_titles = set()
        unique_refs = []
        for ref in references:
            title = ref.get('original_title', ref.get('title', ''))
            if title not in seen_titles:
                unique_refs.append(ref)
                seen_titles.add(title)
        
        for idx, ref in enumerate(unique_refs, start=1):
            formatted = _format_single_reference(ref, idx, ref_format)
            report_lines.append(formatted)
            report_lines.append("\n")
    else:
        report_lines.append("*Nenhuma referência externa foi utilizada neste documento.*\n")

    return "\n".join(report_lines)

def _format_single_reference(ref: dict, idx: int, ref_format: str = "IEEE") -> str:
    title = ref.get('original_title', ref.get('title', 'Untitled'))
    source = ref.get('source', 'Unknown Source')
    year = ref.get('year', '')
    authors = ref.get('authors', None)
    
    if ref_format.upper() == "IEEE":
        author_str = ""
        if authors:
            if isinstance(authors, list):
                author_str = ", ".join(authors)
            else:
                author_str = str(authors)
            author_str += ", "
        return f"[{idx}] {author_str}\"{title}\", {source}, {year}."
    else:
        author_str = ""
        if authors:
            if isinstance(authors, list):
                author_str = ", ".join(authors)
            else:
                author_str = str(authors)
            author_str += ". "
        return f"{author_str}{title}. {source}, {year}."

def save_report(report_id: str, content_md: str, references: Optional[List[dict]] = None):
    """Salva o relatório em disco"""
    reports_dir = Path("reports")
    reports_dir.mkdir(parents=True, exist_ok=True)
    
    md_path = reports_dir / f"{report_id}.md"
    docx_path = reports_dir / f"{report_id}.docx"
    
    # Salvar Markdown
    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content_md)
    
    # Converter para DOCX
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        lines = content_md.splitlines()
        for line in lines:
            line = line.rstrip()
            if line.startswith("# "):
                hdr = doc.add_heading(level=1)
                hdr_run = hdr.add_run(line[2:].strip())
                hdr_run.font.size = Pt(16)
            elif line.startswith("## "):
                hdr = doc.add_heading(level=2)
                hdr_run = hdr.add_run(line[3:].strip())
                hdr_run.font.size = Pt(14)
            elif line.strip() in ["", "---"]:
                continue
            else:
                p = doc.add_paragraph(line)
                p_format = p.paragraph_format
                p_format.space_after = Pt(6)

        doc.save(docx_path)
        logger.info(f"Relatório salvo: {md_path} e {docx_path}")
    except Exception as e:
        logger.error(f"Erro ao salvar DOCX: {e}")